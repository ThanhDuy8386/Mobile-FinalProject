from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\mobi-dev\mobi-project")
OUT = ROOT / "EZ_Finance_Mobile_Report.docx"


PRIMARY_BLUE = "1569FF"
LIGHT_BLUE = "EAF2FB"
LIGHT_GRAY = "F5F7FA"
MID_GRAY = "D9E1EA"
TEXT_GRAY = RGBColor(80, 80, 80)


def set_run_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        if isinstance(color, str):
            run.font.color.rgb = RGBColor.from_string(color)
        else:
            run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=10, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, size=9, color=RGBColor(0, 0, 0))
        set_cell_shading(hdr[idx], LIGHT_BLUE)
        if widths:
            hdr[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=9)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=10)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_sep)
    r._r.append(text)
    r._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    title.font.size = Pt(14)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)

    for style_name, size, center in [
        ("Heading 1", 14, True),
        ("Heading 2", 13, False),
        ("Heading 3", 12.5, False),
        ("Heading 4", 12, False),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        if center:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(6)

    add_page_number_footer(doc)


def add_centered(doc, text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_title(doc, text):
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 13 if level == 2 else 12.5, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True)
    return p


def add_placeholder(doc, title, suggestion):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=260, bottom=260, start=260, end=260)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True, color=RGBColor.from_string(PRIMARY_BLUE))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(suggestion)
    set_run_font(run2, size=10, color=TEXT_GRAY)
    doc.add_paragraph()


def add_cover(doc):
    for text in [
        "EASTERN INTERNATIONAL UNIVERSITY",
        "SCHOOL OF COMPUTING AND INFORMATION TECHNOLOGY",
        "DEPARTMENT OF SOFTWARE ENGINEERING",
    ]:
        add_centered(doc, text, size=13, bold=True, space_after=2)

    doc.add_paragraph()
    doc.add_paragraph()
    add_centered(doc, "MOBILE APPLICATION DEVELOPMENT REPORT", size=16, bold=True, space_after=18)
    add_centered(doc, "EZ FINANCE MOBILE APPLICATION", size=22, bold=True, space_after=28)

    add_centered(doc, "Student(s)", size=12, bold=True, space_after=4)
    for student in [
        "Nguyen Thanh Duy - 2231200009",
        "Nguyen Gia Huy - 2231200078",
        "Nguyen Phat Tai - 2231200025",
    ]:
        add_centered(doc, student, size=12, space_after=2)

    doc.add_paragraph()
    add_centered(doc, "Supervisor(s)", size=12, bold=True, space_after=4)
    add_centered(doc, "Ung Van Giau", size=12, space_after=36)
    add_centered(doc, "Ho Chi Minh City, August 2026", size=12, bold=False, space_after=0)
    doc.add_page_break()


def add_front_matter(doc):
    add_title(doc, "ABSTRACT")
    for text in [
        "EZ Finance is a personal finance mobile application designed to help users record income and expense transactions, organize categories, create monthly budgets, and review financial summaries from a mobile phone. The application focuses on daily usability: users should be able to open the app, check their current balance, add a transaction quickly, and understand whether their spending is still inside the planned budget.",
        "The mobile application is implemented with React Native. It uses stack navigation for detailed flows and bottom tab navigation for the three main areas: Home, Transaction, and Profile. The interface follows a simple card-based layout with clear input forms, color cues for income and expense, floating action buttons for creation flows, dropdown controls for category selection, and date pickers for transaction filtering and entry.",
        "The backend is implemented with Node.js, TypeScript, Express, TypeORM, and MySQL. It provides authentication, user profile management, category CRUD, transaction CRUD, budget CRUD, and report endpoints. The mobile app stores the JWT token in AsyncStorage and sends the token through the Authorization header when calling protected APIs.",
        "Because this report is prepared for a mobile course, the main emphasis is placed on the frontend screens, navigation structure, user interaction design, and API integration from the React Native app. Backend details are summarized only to explain the database structure and the services consumed by the mobile interface.",
    ]:
        add_body(doc, text)
    doc.add_page_break()

    add_title(doc, "ACKNOWLEDGEMENT")
    for text in [
        "We would like to express our sincere gratitude to our supervisor for guiding us during the development of this mobile application. The feedback helped us improve both the technical implementation and the way we present the user interface in a clear, practical, and user-oriented manner.",
        "We also appreciate the support from the School of Computing and Information Technology for providing the learning environment and foundation needed to build a full-stack mobile project. The project allowed us to practice React Native development, REST API integration, data modeling, authentication, and mobile usability design.",
        "Finally, we would like to thank all team members for their contribution to the implementation, testing, and documentation of EZ Finance. The project was completed through collaboration across frontend screens, backend APIs, and report preparation.",
    ]:
        add_body(doc, text)
    doc.add_page_break()

    add_title(doc, "TABLE OF CONTENTS")
    toc_entries = [
        "ABSTRACT",
        "ACKNOWLEDGEMENT",
        "LIST OF FIGURES",
        "LIST OF TABLES",
        "LIST OF ABBREVIATIONS",
        "CHAPTER 1. INTRODUCTION",
        "CHAPTER 2. TECHNOLOGY",
        "CHAPTER 3. APPLICATION ANALYSIS, DESIGN AND IMPLEMENTATION",
        "CHAPTER 4. RESULT AND DISCUSSION",
        "CHAPTER 5. CONCLUSION AND FUTURE WORKS",
        "REFERENCES",
    ]
    for item in toc_entries:
        add_body(doc, item)
    doc.add_page_break()

    add_title(doc, "LIST OF FIGURES")
    figures = [
        "Figure 1. Mobile user use case diagram",
        "Figure 2. Entity Relationship Diagram",
        "Figure 3. Backend project structure",
        "Figure 4. Frontend project structure",
        "Figure 5. Mobile navigation structure",
        "Figure 6. Login screen",
        "Figure 7. Register screen",
        "Figure 8. Home dashboard screen",
        "Figure 9. Monthly report screen",
        "Figure 10. Transaction list screen",
        "Figure 11. Transaction filter screen",
        "Figure 12. Add transaction screen",
        "Figure 13. Transaction detail screen",
        "Figure 14. Category list screen",
        "Figure 15. Add/Edit category screen",
        "Figure 16. Budget list screen",
        "Figure 17. Budget detail screen",
        "Figure 18. Add/Edit budget screen",
        "Figure 19. User profile screen",
        "Figure 20. Profile update and change password screens",
    ]
    for figure in figures:
        add_body(doc, figure)
    doc.add_page_break()

    add_title(doc, "LIST OF TABLES")
    tables = [
        "Table 1. List of abbreviations",
        "Table 2. Main mobile technologies",
        "Table 3. Backend database overview",
        "Table 4. API endpoints used by the mobile application",
        "Table 5. Frontend screen organization",
        "Table 6. Main UI results",
    ]
    for table in tables:
        add_body(doc, table)
    doc.add_page_break()

    add_title(doc, "LIST OF ABBREVIATIONS")
    add_table(
        doc,
        ["No.", "Term", "Meaning"],
        [
            [1, "API", "Application Programming Interface"],
            [2, "CRUD", "Create, Read, Update, Delete"],
            [3, "DTO", "Data Transfer Object"],
            [4, "JWT", "JSON Web Token"],
            [5, "ORM", "Object Relational Mapping"],
            [6, "REST", "Representational State Transfer"],
            [7, "UI", "User Interface"],
            [8, "UX", "User Experience"],
        ],
        [0.5, 1.1, 4.5],
    )
    add_caption(doc, "Table 1. List of abbreviations")
    doc.add_page_break()


def add_chapter_1(doc):
    add_heading(doc, "INTRODUCTION", 1)

    add_heading(doc, "Motivation", 2)
    for text in [
        "Personal finance management is a common problem for students and young workers. Many people remember large expenses, but they often forget small daily payments such as food, transport, coffee, or online purchases. These small expenses make it difficult to understand where money is going and whether the user is still following a monthly plan.",
        "A mobile application is suitable for this problem because financial records are created in daily life. Users can add a transaction immediately after spending money, check the monthly dashboard before making a purchase, and review their budget progress at any time. Compared with a web-only system, a mobile app is closer to the user's daily behavior and can provide faster interaction for repeated tasks.",
        "EZ Finance is built to support these daily activities. The application is not only a CRUD interface for financial data; it also organizes the information into dashboard summaries, monthly reports, recent transaction lists, and budget progress indicators. For the mobile course, the project demonstrates how a React Native interface can connect to a REST API and present database information through practical mobile screens.",
    ]:
        add_body(doc, text)

    add_heading(doc, "Project Objectives", 2)
    for item in [
        "Build a mobile authentication flow that supports register, login, token storage, logout, profile update, and password change.",
        "Design a dashboard screen that summarizes total income, total expense, current monthly balance, all-time balance, recent transactions, and budget status.",
        "Provide transaction management screens for listing, searching, filtering, creating, viewing details, updating, and deleting income or expense transactions.",
        "Provide category management screens so users can organize transactions by income and expense categories with icon and color metadata.",
        "Provide budget management screens so users can create monthly expense limits, view spent amount, remaining amount, percentage used, and related expense transactions.",
        "Integrate the mobile app with backend APIs using JWT Bearer authentication and consistent JSON response handling.",
        "Prepare a report focused on mobile UI design, screen organization, interaction flow, and API usage.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Project Scope", 2)
    for text in [
        "The scope of this report is the EZ Finance mobile application. The frontend is the main focus because the course topic is mobile development. The report explains the user-facing screens, navigation hierarchy, visual layout decisions, state management, and API calls made from the React Native source code.",
        "The backend is included only as supporting context. It is described through the database tables, authentication mechanism, business rules, and REST endpoints consumed by the mobile app. Detailed backend implementation such as controller code, service classes, Swagger schema, and migration logic is summarized instead of being analyzed in full depth.",
    ]:
        add_body(doc, text)

    add_heading(doc, "Technology Overview", 2)
    add_body(
        doc,
        "The mobile frontend is built with React Native, React Navigation, AsyncStorage, DateTimePicker, Safe Area Context, and Ionicons. The backend uses Express, TypeScript, TypeORM, MySQL, JWT, bcrypt, class-validator, CORS, Helmet, Morgan, and Swagger UI. The mobile app communicates with the backend through REST endpoints using the fetch API."
    )

    add_heading(doc, "Report Structure", 2)
    for item in [
        "Chapter 1 introduces the motivation, objectives, scope, and technologies of the project.",
        "Chapter 2 explains the technologies used in the mobile application and briefly introduces the backend stack.",
        "Chapter 3 presents system analysis, database overview, API integration, mobile navigation, and frontend implementation.",
        "Chapter 4 focuses on the implemented mobile screens and discusses the result of each UI flow.",
        "Chapter 5 concludes the project and suggests future improvements.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_chapter_2(doc):
    add_heading(doc, "TECHNOLOGY", 1)

    add_heading(doc, "React Native", 2)
    for text in [
        "React Native is the main framework used to build the mobile interface of EZ Finance. It allows the project to define screens with JavaScript and React components while still rendering native mobile UI elements. This is suitable for a course project because one codebase can target Android and iOS structures while keeping the implementation familiar to React developers.",
        "In this project, React Native components such as View, Text, TextInput, TouchableOpacity, ScrollView, FlatList, ActivityIndicator, Alert, and StyleSheet are used to build the interface. The screens use cards, form inputs, dropdown-like controls, colored status text, progress bars, and floating action buttons to make financial data easy to scan on a small screen.",
    ]:
        add_body(doc, text)

    add_heading(doc, "React Navigation", 2)
    for text in [
        "React Navigation organizes the application into a root stack, bottom tabs, and nested stacks. The root stack contains Login, Register, ForgotPassword, and MainTabs. After successful authentication, the user is moved to MainTabs. The main tabs are Home, Transaction, and Profile.",
        "Each tab has its own navigation responsibility. Home contains the dashboard and monthly report. Transaction contains the transaction list, detail, add, edit, and filter screens. Profile contains user profile features as well as category and budget management. This structure keeps common workflows short while still allowing detailed screens to be pushed when needed.",
    ]:
        add_body(doc, text)

    add_heading(doc, "AsyncStorage", 2)
    add_body(
        doc,
        "The mobile app uses AsyncStorage to store the JWT token after login or registration. Protected requests read the token and send it in the Authorization header as Bearer token. Logout removes the token from local storage and resets navigation back to the Login screen."
    )

    add_heading(doc, "Mobile UI Support Libraries", 2)
    add_table(
        doc,
        ["Technology", "Used in project", "Purpose"],
        [
            ["@react-navigation/native", "Yes", "Navigation container for the app."],
            ["@react-navigation/native-stack", "Yes", "Stack flows for auth, profile, transaction, and home detail screens."],
            ["@react-navigation/bottom-tabs", "Yes", "Bottom tab navigation for Home, Transaction, and Profile."],
            ["@react-native-async-storage/async-storage", "Yes", "Stores JWT token locally."],
            ["@react-native-community/datetimepicker", "Yes", "Date selection for transaction forms and advanced filters."],
            ["react-native-vector-icons", "Yes", "Ionicons for tabs, buttons, password visibility, date, search, edit, delete, and logout."],
            ["react-native-safe-area-context", "Yes", "Safe area padding for screens with custom layouts."],
            ["react-native-svg", "Installed", "Available for future SVG charts or custom visualizations."],
        ],
        [1.7, 0.9, 3.5],
    )
    add_caption(doc, "Table 2. Main mobile technologies")

    add_heading(doc, "Backend Stack Overview", 2)
    add_body(
        doc,
        "The backend is a REST API built with Node.js, TypeScript, Express, TypeORM, and MySQL. Express handles HTTP routing, TypeORM maps TypeScript entities to MySQL tables, class-validator validates DTO request bodies, bcrypt hashes passwords, and JWT protects user-specific resources. The backend also exposes Swagger UI for API testing and documentation."
    )

    add_heading(doc, "REST API Communication", 2)
    for text in [
        "The frontend uses the native fetch API to call backend endpoints. In the Android emulator environment, the source code points to http://10.0.2.2:5001/api. The backend default port in the configuration is 5000, so the development environment must either run the backend on port 5001 or update the mobile base URL to match the actual server port.",
        "Most endpoints are protected. The mobile app retrieves the token from AsyncStorage, includes it in the Authorization header, and then reads the common response format: success, message, data, and optional pagination. This pattern appears in profile, category, transaction, budget, and report screens.",
    ]:
        add_body(doc, text)
    doc.add_page_break()


def add_database_section(doc):
    add_heading(doc, "Database", 2)
    add_heading(doc, "Entity Relationship Diagram", 3)
    add_placeholder(
        doc,
        "Image placeholder",
        "Suggested image: ERD showing users related to categories, transactions, and budgets. Categories connect to transactions and budgets.",
    )
    add_caption(doc, "Figure 2. Entity Relationship Diagram")

    add_heading(doc, "Database Analysis", 3)
    add_body(
        doc,
        "The backend database contains four main tables. The users table stores account information and authentication data. The categories table stores user-defined income and expense categories. The transactions table stores each income or expense record. The budgets table stores a monthly spending limit for a selected expense category. All financial records are connected to a user, so each user can only access their own data."
    )
    add_table(
        doc,
        ["Table", "Main fields", "Relationships", "Role in app"],
        [
            [
                "users",
                "id, fullName, email, passwordHash, createdAt, updatedAt",
                "One user has many categories, transactions, and budgets.",
                "Stores user accounts. Passwords are stored as bcrypt hashes and are not returned to the mobile app.",
            ],
            [
                "categories",
                "id, name, type, icon, color, userId, createdAt, updatedAt",
                "Belongs to one user. Has many transactions and budgets.",
                "Classifies transactions as INCOME or EXPENSE. Expense categories can be used to create budgets.",
            ],
            [
                "transactions",
                "id, title, amount, type, transactionDate, note, userId, categoryId, createdAt, updatedAt",
                "Belongs to one user and one category.",
                "Stores the user's financial activities and feeds dashboard, monthly report, category report, and budget progress.",
            ],
            [
                "budgets",
                "id, limitAmount, month, year, userId, categoryId, createdAt, updatedAt",
                "Belongs to one user and one expense category.",
                "Stores spending limits. Spent amount, remaining amount, percentage, and exceeded status are calculated from transactions.",
            ],
        ],
        [0.9, 1.7, 1.5, 2.0],
    )
    add_caption(doc, "Table 3. Backend database overview")

    add_heading(doc, "Important Backend Rules", 3)
    for item in [
        "Users can only access categories, transactions, and budgets that belong to their own user ID.",
        "Registering a user automatically creates default income and expense categories.",
        "Category names must be unique for the same user and same category type.",
        "A category cannot be deleted if it is already used by transactions or budgets.",
        "Transaction amount must be greater than zero and transaction type must match the selected category type.",
        "Only expense categories can have budgets.",
        "A user cannot create duplicate budgets for the same category, month, and year.",
        "Budget progress is calculated from expense transactions instead of being stored directly.",
    ]:
        add_bullet(doc, item)


def add_api_section(doc):
    add_heading(doc, "APIs Used By The Mobile Application", 2)
    add_body(
        doc,
        "The following table summarizes the backend APIs that are called from the React Native source code. These endpoints are the bridge between the mobile UI and the MySQL database."
    )
    add_table(
        doc,
        ["Mobile screen", "Method and endpoint", "Purpose"],
        [
            ["LoginScreen", "POST /api/auth/login", "Authenticate user, receive JWT token, store token, and enter MainTabs."],
            ["RegisterScreen", "POST /api/auth/register", "Create account, receive JWT token, and enter MainTabs."],
            ["UserProfileScreen", "GET /api/users/profile", "Load current user's full name and email."],
            ["EditProfileScreen", "PUT /api/users/profile", "Update full name and email."],
            ["ChangePasswordScreen", "PUT /api/users/change-password", "Change password after validating current and new password."],
            ["CategoryListScreen", "GET /api/categories?type=INCOME or EXPENSE", "Load categories by selected type."],
            ["AddCategoryScreen", "POST /api/categories", "Create a new income or expense category."],
            ["EditCategoryScreen", "PUT /api/categories/:id", "Update category name, type, icon, and color."],
            ["CategoryListScreen", "DELETE /api/categories/:id", "Delete a category if it is not referenced by transactions or budgets."],
            ["TransactionScreen", "GET /api/transactions?...query", "Load transaction list with pagination, type filter, keyword, category, date, month, year, and sorting."],
            ["TransactionDetailScreen", "GET /api/transactions/:id", "Refresh and display transaction detail."],
            ["AddTransactionScreen", "GET /api/categories?type=...", "Load matching categories before creating transaction."],
            ["AddTransactionScreen", "POST /api/transactions", "Create income or expense transaction."],
            ["EditTransactionScreen", "GET /api/categories?type=...", "Reload category options when transaction type changes."],
            ["EditTransactionScreen", "PUT /api/transactions/:id", "Update title, amount, type, category, date, and note."],
            ["TransactionDetailScreen", "DELETE /api/transactions/:id", "Delete selected transaction."],
            ["HomeDashboardScreen", "GET /api/reports/dashboard", "Load totals, recent transactions, and budget summary for dashboard."],
            ["MonthlyReportScreen", "GET /api/reports/monthly?year=2026", "Load income, expense, and balance for each month in the selected year."],
            ["BudgetListScreen", "GET /api/budgets?month=...&year=...", "Load budgets for selected month and year."],
            ["AddBudgetScreen", "GET /api/categories?type=EXPENSE", "Load expense categories for budget creation."],
            ["AddBudgetScreen", "POST /api/budgets", "Create a budget for an expense category."],
            ["BudgetDetailScreen", "GET /api/budgets/:id", "Load budget detail with related expense transactions."],
            ["EditBudgetScreen", "GET /api/categories?type=EXPENSE", "Load available expense categories for editing budget."],
            ["EditBudgetScreen", "PUT /api/budgets/:id", "Update budget category, limit amount, month, and year."],
        ],
        [1.4, 1.9, 2.8],
    )
    add_caption(doc, "Table 4. API endpoints used by the mobile application")

    add_body(
        doc,
        "The backend also provides GET /api/health, GET /api/auth/me, DELETE /api/budgets/:id, GET /api/reports/expenses-by-category, and GET /api/reports/income-by-category. These endpoints are available in the backend but are not directly used by the current mobile screens, except that they may be useful for future health checks, token validation, budget deletion, or chart visualizations."
    )


def add_chapter_3(doc):
    add_heading(doc, "APPLICATION ANALYSIS, DESIGN AND IMPLEMENTATION", 1)

    add_heading(doc, "System Description", 2)
    for text in [
        "EZ Finance is a mobile personal finance management system. A user creates an account or logs in with an existing account, then uses the app to manage daily financial data. The main app area is divided into Home, Transaction, and Profile tabs.",
        "The Home tab gives a quick overview of the user's financial status. The Transaction tab provides the main workflow for recording and reviewing money movement. The Profile tab groups personal account actions and management screens for categories and budgets.",
    ]:
        add_body(doc, text)

    add_heading(doc, "Mobile User Use Cases", 2)
    for item in [
        "Register a new account and start using the app immediately after token creation.",
        "Log in with email and password, then keep the session by storing JWT token locally.",
        "View dashboard totals, recent transactions, and budget summary.",
        "View monthly report for all twelve months in a year.",
        "Create, search, filter, view, edit, and delete transactions.",
        "Create, view, edit, and delete categories by income or expense type.",
        "Create, view, and edit monthly budgets for expense categories.",
        "View profile, edit profile information, change password, and log out.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Interface Design Principles", 2)
    for item in [
        "Keep the main navigation simple with three bottom tabs: Home, Transaction, and Profile.",
        "Use card layouts for financial summaries, transaction rows, profile information, and budget status.",
        "Use green for income or normal budget status and red for expense, deletion, or exceeded budget status.",
        "Use floating add buttons in list screens so creation actions are always easy to reach.",
        "Use dropdown controls for type/category selection and DateTimePicker for date input instead of free-form date typing where possible.",
        "Show loading indicators, retry states, empty states, and validation messages to make network-based mobile flows clearer.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Use Case Diagram", 2)
    add_placeholder(
        doc,
        "Image placeholder",
        "Suggested image: Mobile user use case diagram with Authentication, Dashboard, Transaction, Category, Budget, Report, and Profile modules.",
    )
    add_caption(doc, "Figure 1. Mobile user use case diagram")

    add_database_section(doc)
    add_api_section(doc)

    add_heading(doc, "Backend Project Structure", 2)
    add_placeholder(
        doc,
        "Image placeholder",
        "Suggested image: screenshot of ez-finance-api/src showing config, controllers, database, dtos, entities, middleware, routes, services, utils.",
    )
    add_caption(doc, "Figure 3. Backend project structure")
    add_body(
        doc,
        "The backend follows a layered structure. Routes define endpoint paths and middleware. Controllers receive HTTP requests and return formatted responses. Services contain business logic such as validation beyond DTOs, ownership checks, unique constraints, and calculations. Entities define database tables. DTOs define input validation rules. Middleware handles authentication, validation, not-found responses, and error responses."
    )

    add_heading(doc, "Frontend Project Structure", 2)
    add_placeholder(
        doc,
        "Image placeholder",
        "Suggested image: screenshot of ez-finance-app/EzFinanceApp/src showing all screen files and App.js navigation.",
    )
    add_caption(doc, "Figure 4. Frontend project structure")
    add_table(
        doc,
        ["Area", "Screens", "Main responsibility"],
        [
            ["Authentication", "LoginScreen, RegisterScreen, ForgotPasswordScreen", "Entry flow, account creation, password visibility controls, and token storage after successful auth."],
            ["Home", "HomeDashboardScreen, MonthlyReportScreen", "Financial overview, report summary, recent transactions, and budget status."],
            ["Transaction", "TransactionScreen, TransactionFilterScreen, TransactionDetailScreen, AddTransactionScreen, EditTransactionScreen", "Core transaction workflow including list, filter, detail, create, update, and delete."],
            ["Category", "CategoryListScreen, AddCategoryScreen, EditCategoryScreen", "Manage income and expense categories with type, icon, and color."],
            ["Budget", "BudgetListScreen, BudgetDetailScreen, AddBudgetScreen, EditBudgetScreen", "Manage monthly expense budgets and show budget progress."],
            ["Profile", "UserProfileScreen, EditProfileScreen, ChangePasswordScreen", "Account profile, personal settings, password change, and logout."],
        ],
        [1.2, 2.5, 2.4],
    )
    add_caption(doc, "Table 5. Frontend screen organization")

    add_heading(doc, "Navigation Structure", 2)
    add_placeholder(
        doc,
        "Image placeholder",
        "Suggested image: navigation map showing RootStack -> Login/Register/ForgotPassword/MainTabs; MainTabs -> HomeStack, TransactionStack, ProfileStack.",
    )
    add_caption(doc, "Figure 5. Mobile navigation structure")
    for text in [
        "The root stack starts at Login. Login and Register both call authentication APIs and then replace the route with MainTabs, which prevents the user from returning to the auth screen by pressing back. MainTabs contains the three primary sections of the mobile app.",
        "HomeStack contains HomeDashboard and MonthlyReport. TransactionStack contains TransactionList, TransactionDetail, AddTransaction, EditTransaction, and TransactionFilter. ProfileStack contains UserProfile, EditProfile, ChangePassword, CategoryList, AddCategory, EditCategory, BudgetList, BudgetDetail, AddBudget, and EditBudget.",
    ]:
        add_body(doc, text)

    add_heading(doc, "Frontend Implementation", 2)
    add_heading(doc, "State Management", 3)
    add_body(
        doc,
        "The current project uses React local state through useState and screen lifecycle hooks such as useEffect and useFocusEffect. Because the app is small and screen-based, local state is enough for form data, list data, loading state, error messages, dropdown visibility, selected filters, and selected dates."
    )
    add_heading(doc, "Authentication Handling", 3)
    add_body(
        doc,
        "After login or registration, the token returned by the backend is stored in AsyncStorage under the key token. Protected screens read this token before sending requests. Logout removes the token and resets the navigation stack back to Login."
    )
    add_heading(doc, "Form Validation", 3)
    add_body(
        doc,
        "Several forms perform client-side validation before calling the API. Register validates full name length, email format, and password length. Transaction forms validate title, amount, category, and date. Budget forms validate category, positive limit amount, valid month, and year. ChangePassword validates missing fields, minimum length, and whether the new password differs from the current password."
    )
    add_heading(doc, "Loading and Error Feedback", 3)
    add_body(
        doc,
        "The app uses ActivityIndicator for loading states, Alert dialogs for action results or errors, retry buttons on some screens, and empty-state text for lists with no data. This feedback is important in a mobile app because most screens depend on network requests."
    )
    doc.add_page_break()


def add_screen_result(doc, heading, figure_caption, suggestion, paragraphs):
    add_heading(doc, heading, 3)
    for text in paragraphs:
        add_body(doc, text)
    add_placeholder(doc, "Screenshot placeholder", suggestion)
    add_caption(doc, figure_caption)


def add_chapter_4(doc):
    add_heading(doc, "RESULT AND DISCUSSION", 1)

    add_heading(doc, "Installation Environment", 2)
    add_heading(doc, "Development Environment", 3)
    for item in [
        "Frontend: React Native CLI project located at ez-finance-app/EzFinanceApp.",
        "Backend: Node.js and TypeScript API located at ez-finance-api.",
        "Database: MySQL database named ez_finance.",
        "Android emulator API access: http://10.0.2.2:5001/api in the mobile source code.",
        "Backend default configuration: PORT fallback is 5000, so development port must be aligned with the mobile app before running a demo.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Framework and Dependencies", 3)
    add_body(
        doc,
        "The frontend package includes React Native 0.86.2, React 19.2.3, React Navigation, AsyncStorage, DateTimePicker, Safe Area Context, React Native Screens, React Native SVG, Vector Icons, Jest, ESLint, Prettier, and TypeScript configuration. The backend package includes Express, TypeORM, MySQL2, JWT, bcrypt, class-validator, class-transformer, dotenv, CORS, Helmet, Morgan, and Swagger UI Express."
    )

    add_heading(doc, "Build Configuration", 3)
    for item in [
        "Run the backend with npm run dev from ez-finance-api after creating the MySQL database, running migrations, and optionally seeding demo data.",
        "Run the mobile Metro server with npm start from ez-finance-app/EzFinanceApp.",
        "Run Android with npm run android or iOS with npm run ios depending on the development environment.",
        "For Android emulator demo, make sure the backend is reachable through 10.0.2.2 and the port matches the mobile source code.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Results and Discussions", 2)
    add_body(
        doc,
        "This section focuses on the implemented mobile screens. Each subsection describes the purpose, layout, main controls, and API integration of the screen. The screenshot boxes are left as placeholders so final images can be captured from the emulator and inserted later."
    )

    add_screen_result(
        doc,
        "Login Screen",
        "Figure 6. Login screen",
        "Capture the login screen with email field, password field, show/hide password icon, Forgot password link, Login button, and Register link.",
        [
            "The Login screen is the first screen of the application. It contains a centered title, email input, password input, password visibility icon, Forgot password link, and Login button.",
            "When the user presses Login, the screen calls POST /api/auth/login. If the response is successful, the JWT token is stored in AsyncStorage and the navigation route is replaced with MainTabs. If login fails, an Alert displays the backend error message.",
        ],
    )
    add_screen_result(
        doc,
        "Register Screen",
        "Figure 7. Register screen",
        "Capture the register screen with full name, email, password, show/hide password icon, Register button, and Login link.",
        [
            "The Register screen allows a new user to create an account. It validates full name length, email format, and password length before sending a request to the backend.",
            "After successful registration, the backend automatically creates default categories for the new user, returns a token, and the mobile app stores that token before entering the main application.",
        ],
    )
    add_screen_result(
        doc,
        "Home Dashboard Screen",
        "Figure 8. Home dashboard screen",
        "Capture the Home tab showing total income, total expense, balance, all-time balance, recent transactions, and budget summary.",
        [
            "The Home dashboard is the main financial overview screen. It displays four summary cards: Total Income, Total Expense, Balance, and All Time Balance. These values come from GET /api/reports/dashboard.",
            "Below the summary cards, the screen shows a View Monthly Report button, a Recent Transactions section, and a Budget Summary table. Income and expense are visually distinguished with green and red text. Budget percentage is shown in red when exceeded.",
        ],
    )
    add_screen_result(
        doc,
        "Monthly Report Screen",
        "Figure 9. Monthly report screen",
        "Capture the monthly report list showing month number, total income, total expense, and balance.",
        [
            "The Monthly Report screen uses a FlatList to display twelve monthly cards. Each card shows the month number, total income, total expense, and balance.",
            "The current source calls GET /api/reports/monthly?year=2026. In a future version, the year should be selectable from the interface instead of being hard-coded.",
        ],
    )
    add_screen_result(
        doc,
        "Transaction List Screen",
        "Figure 10. Transaction list screen",
        "Capture the Transaction tab with ALL/INCOME/EXPENSE filter pills, search bar, filter icon, transaction cards, and floating add button.",
        [
            "The Transaction List screen is the central workflow of the app. It displays transactions in cards with title, signed amount, category, type, date, and note. Expenses are shown with a minus sign and red amount, while income is shown with a plus sign and green amount.",
            "The screen supports quick type filters, keyword search, pull-to-refresh, pagination through onEndReached, loading indicators, retry messages, and an empty state. Pressing a transaction opens its detail screen, while the floating button opens the add transaction form.",
        ],
    )
    add_screen_result(
        doc,
        "Transaction Filter Screen",
        "Figure 11. Transaction filter screen",
        "Capture the advanced filter screen with type, category, month, year, start date, end date, keyword, sort by, sort order, and Apply Filter button.",
        [
            "The Transaction Filter screen provides advanced filtering controls. The user can filter by type, category, month, year, date range, and keyword. The user can also choose the sort field and sort order.",
            "When Apply Filter is pressed, the selected filters are passed back to TransactionList through navigation params. TransactionScreen then builds the query string and reloads GET /api/transactions with the selected parameters.",
        ],
    )
    add_screen_result(
        doc,
        "Add Transaction Screen",
        "Figure 12. Add transaction screen",
        "Capture the add transaction form with title, amount, type dropdown, category dropdown, date picker, note input, and Add button.",
        [
            "The Add Transaction screen is designed as a form inside a bordered card. It includes title, amount, type, category, transaction date, note, and save action.",
            "When the selected type changes, the screen reloads categories using GET /api/categories?type=INCOME or EXPENSE. This prevents the user from selecting an expense category for an income transaction or the opposite. The form then sends POST /api/transactions.",
        ],
    )
    add_screen_result(
        doc,
        "Transaction Detail Screen",
        "Figure 13. Transaction detail screen",
        "Capture the transaction detail card with title, amount, type, category, date, note, Delete button, and Edit button.",
        [
            "The Transaction Detail screen shows one transaction in a simple key-value layout. It refreshes the record by calling GET /api/transactions/:id when the screen is focused.",
            "The screen includes two main actions: Delete and Edit. Delete asks for confirmation and then calls DELETE /api/transactions/:id. Edit navigates to EditTransaction with the current transaction data.",
        ],
    )
    add_screen_result(
        doc,
        "Category List Screen",
        "Figure 14. Category list screen",
        "Capture category list with INCOME/EXPENSE toggle, category rows, edit/delete actions, and floating add button.",
        [
            "The Category List screen separates categories by INCOME and EXPENSE filters. Each row shows a color circle, category name, type, edit action, and delete action.",
            "The screen reloads when focused, so updates from add or edit screens are reflected when the user returns. Delete uses a confirmation dialog and calls DELETE /api/categories/:id. The backend prevents deletion when the category is already used by transactions or budgets.",
        ],
    )
    add_screen_result(
        doc,
        "Add and Edit Category Screens",
        "Figure 15. Add/Edit category screen",
        "Capture the category form showing name, INCOME/EXPENSE segmented type buttons, icon dropdown, color dropdown, and save button.",
        [
            "The Add Category and Edit Category screens share a similar structure. The user can enter a category name, choose a type, choose an icon text value, and choose a color from predefined color options.",
            "The Add screen sends POST /api/categories, while the Edit screen sends PUT /api/categories/:id. Both screens show Alert messages for success or backend validation errors.",
        ],
    )
    add_screen_result(
        doc,
        "Budget List Screen",
        "Figure 16. Budget list screen",
        "Capture budget list with month selector, year controls, budget cards, progress bar, exceeded status, and floating add button.",
        [
            "The Budget List screen lets the user choose a month and year. It calls GET /api/budgets?month=...&year=... and displays a card for each budget.",
            "Each budget card shows category, month, year, limit amount, spent amount, remaining amount, percentage, exceeded status, and a progress bar. The progress bar uses green for normal progress and red when the budget is exceeded.",
        ],
    )
    add_screen_result(
        doc,
        "Budget Detail Screen",
        "Figure 17. Budget detail screen",
        "Capture budget detail with limit/spent/remaining/percentage/month/year and related expense transactions.",
        [
            "The Budget Detail screen displays detailed budget progress and a list of related expense transactions for the selected budget category, month, and year.",
            "The screen calls GET /api/budgets/:id. The backend calculates spentAmount, remainingAmount, percentage, and isExceeded, then returns expense transactions related to that budget. The screen also provides an Edit Budget button.",
        ],
    )
    add_screen_result(
        doc,
        "Add and Edit Budget Screens",
        "Figure 18. Add/Edit budget screen",
        "Capture the budget form with expense category dropdown, limit amount, month, year, and save button.",
        [
            "The Add Budget and Edit Budget screens allow the user to select an expense category, enter a limit amount, month, and year. They load available expense categories through GET /api/categories?type=EXPENSE.",
            "The Add screen calls POST /api/budgets and the Edit screen calls PUT /api/budgets/:id. The UI validates missing category, invalid limit amount, and invalid month before sending the request.",
        ],
    )
    add_screen_result(
        doc,
        "User Profile Screen",
        "Figure 19. User profile screen",
        "Capture the Profile tab with avatar, full name, email, menu list, and Logout button.",
        [
            "The User Profile screen displays user information in a profile card and provides menu entries for Category List, Budget List, Edit Profile, and Change Password.",
            "The screen loads profile data through GET /api/users/profile when it is focused. Logout removes the stored token and resets the root navigation stack to Login.",
        ],
    )
    add_screen_result(
        doc,
        "Edit Profile and Change Password Screens",
        "Figure 20. Profile update and change password screens",
        "Capture Edit Profile and Change Password screens. If space is limited, combine both screenshots into one figure.",
        [
            "The Edit Profile screen allows the user to update full name and email with client-side validation before calling PUT /api/users/profile.",
            "The Change Password screen provides current and new password inputs, eye icons to show or hide the password, validation, a loading state, and PUT /api/users/change-password integration.",
        ],
    )

    add_heading(doc, "Discussion", 2)
    for text in [
        "The project successfully implements the main UI flows expected from a personal finance app. The Home tab gives a compact financial overview, the Transaction tab supports a complete transaction lifecycle, and the Profile tab collects account, category, and budget management features.",
        "The strongest part of the mobile implementation is the organization of screens around user tasks. Bottom tabs reduce navigation complexity, while nested stack screens allow users to move from list to detail to edit forms naturally. The use of FlatList, RefreshControl, ActivityIndicator, Alert, and local validation makes the app practical for real mobile usage.",
        "There are also several limitations. The forgot password screen is currently only a UI prototype and does not call a backend reset-password API. The monthly report screen currently uses a fixed year 2026. The API base URL is repeated in several files instead of being centralized in one configuration module. The app can also be improved with charts, dark mode, stronger empty states, and automated UI tests.",
    ]:
        add_body(doc, text)

    add_table(
        doc,
        ["Feature area", "Implemented result", "Discussion"],
        [
            ["Authentication", "Login and register call backend APIs and store JWT token.", "Good basic flow; forgot password still needs backend integration."],
            ["Dashboard", "Shows totals, recent transactions, and budget summary.", "Useful first screen; future charts would improve scanning."],
            ["Transactions", "List, search, filter, add, detail, edit, delete.", "Most complete module in the mobile app."],
            ["Categories", "List by type, create, edit, delete.", "Supports transaction organization and budget setup."],
            ["Budgets", "List by month/year, detail, add, edit, progress bar.", "Shows calculated spending status clearly."],
            ["Profile", "View profile, edit profile, change password, logout.", "Covers core account management."],
        ],
        [1.2, 2.4, 2.5],
    )
    add_caption(doc, "Table 6. Main UI results")
    doc.add_page_break()


def add_chapter_5(doc):
    add_heading(doc, "CONCLUSION AND FUTURE WORKS", 1)

    add_heading(doc, "Conclusion", 2)
    for text in [
        "EZ Finance demonstrates a complete mobile interface for personal finance management. The app allows users to authenticate, view a financial dashboard, manage transactions, manage categories, manage budgets, and update profile information. The project shows how React Native can be used to build practical mobile screens that connect to a REST API and display user-specific data from a backend database.",
        "From a mobile development perspective, the project uses appropriate navigation patterns, reusable screen layouts, native input controls, and clear visual feedback. The Transaction module is the most complete workflow because it includes listing, filtering, detail view, creation, editing, and deletion. The Budget module also provides a meaningful mobile experience by translating backend calculations into progress indicators and clear status text.",
        "The backend supports the mobile app through a simple but useful database design and API set. The main tables are users, categories, transactions, and budgets. The API protects user data with JWT and applies business rules that keep financial data consistent.",
    ]:
        add_body(doc, text)

    add_heading(doc, "Future Works", 2)
    for item in [
        "Centralize the API base URL in one configuration file so the app can switch between emulator, physical device, and production API more easily.",
        "Add a real forgot-password backend flow instead of keeping the current reset-password UI as a prototype.",
        "Add selectable year and month controls to MonthlyReportScreen instead of hard-coding year 2026.",
        "Use react-native-svg or another chart library to visualize monthly income/expense and category distribution.",
        "Add budget deletion from the mobile UI because the backend already supports DELETE /api/budgets/:id.",
        "Add token validation or auto-login using GET /api/auth/me when the app starts.",
        "Improve offline behavior with cached dashboard and transaction data.",
        "Add push notifications for budgets that are close to being exceeded.",
        "Add dark mode and more consistent design tokens for colors, spacing, and typography.",
        "Add automated tests for navigation, form validation, API handling, and important UI states.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    add_title(doc, "REFERENCES")
    refs = [
        "[1] React Native Documentation. https://reactnative.dev/",
        "[2] React Navigation Documentation. https://reactnavigation.org/",
        "[3] AsyncStorage Documentation. https://react-native-async-storage.github.io/async-storage/",
        "[4] Express.js Documentation. https://expressjs.com/",
        "[5] TypeORM Documentation. https://typeorm.io/",
        "[6] MySQL Documentation. https://dev.mysql.com/doc/",
        "[7] JSON Web Token Introduction. https://jwt.io/introduction",
        "[8] EZ Finance project source code, ez-finance-api and ez-finance-app/EzFinanceApp.",
    ]
    for ref in refs:
        add_body(doc, ref)


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)

    doc.core_properties.title = "EZ Finance Mobile Application Report"
    doc.core_properties.subject = "Mobile Application Development"
    doc.core_properties.author = "EZ Finance Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
